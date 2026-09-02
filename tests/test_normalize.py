"""``memoria normalize``: the skip-unchanged normalization run (part 05 §5.4)."""

import base64
import io
import mailbox
import sqlite3
from dataclasses import replace
from email.message import EmailMessage

from docx import Document
from docx.shared import Inches

from memoria.index import INDEX_RELATIVE_PATH, build_index, search
from memoria.manifest import (
    DEFAULT_MANIFEST_RELATIVE_PATH,
    load_converter_pins,
    load_manifest,
)
from memoria.normalize import CONVERTERS, EMAIL_CONVERTER_VERSION, normalize
from memoria.records import (
    NORMALIZED_RELATIVE_PATH,
    read_all,
    real_paragraphs,
    record_to_markdown,
)
from memoria.repository import Repository
from memoria.validate import validate


def _write_raw_file(evidence_root, rel_path, content):
    full = evidence_root / "raw" / rel_path
    full.parent.mkdir(parents=True, exist_ok=True)
    full.write_text(content)
    return full


def _email_message(
    *, from_, to, date, message_id, body, cc=None, in_reply_to=None, attachment=None
):
    message = EmailMessage()
    message["From"] = from_
    message["To"] = to
    if cc:
        message["Cc"] = cc
    message["Date"] = date
    message["Message-ID"] = message_id
    if in_reply_to:
        message["In-Reply-To"] = in_reply_to
    message.set_content(body)
    if attachment is not None:
        filename, content_bytes = attachment
        message.add_attachment(
            content_bytes,
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=filename,
        )
    return message


def _write_mbox(evidence_root, rel_path, messages):
    full = evidence_root / "raw" / rel_path
    full.parent.mkdir(parents=True, exist_ok=True)
    box = mailbox.mbox(str(full))
    for message in messages:
        box.add(message)
    box.flush()
    box.close()
    return full


def _write_raw_binary_file(evidence_root, rel_path, data):
    full = evidence_root / "raw" / rel_path
    full.parent.mkdir(parents=True, exist_ok=True)
    full.write_bytes(data)
    return full


# --- fixture builders (#77) -------------------------------------------------
#
# No PDF-writing library is a project dependency, so both fixtures are built
# by hand: docx via python-docx (test-only, `dev`'s own extra - see
# pyproject.toml), pdf as the minimal object structure pdfplumber needs to
# read one back.


def _make_docx(*, with_image=True):
    """A docx with a heading, a bold/italic run, a bullet list, a table and,
    unless ``with_image=False``, one embedded image."""
    document = Document()
    document.add_heading("A Sample Heading", level=1)

    paragraph = document.add_paragraph("This paragraph has ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(" text.")

    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("Second bullet", style="List Bullet")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Row1 A"
    table.cell(1, 1).text = "Row1 B"

    if with_image:
        # A real, minimal 1x1 PNG - it is the docx's own zip entry
        # (word/media/...) that names the image, never the pixel content,
        # but python-docx still parses it as an image on the way in.
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4"
            "2mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        document.add_picture(io.BytesIO(png), width=Inches(1))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_from_objects(objects):
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode()
        out += obj
        out += b"\nendobj\n"
    xref_offset = len(out)
    count = len(objects) + 1
    out += f"xref\n0 {count}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {count} /Root 1 0 R >>\n".encode()
    out += f"startxref\n{xref_offset}\n%%EOF".encode()
    return bytes(out)


def _make_pdf(pages):
    """A minimal hand-built pdf. Each entry in ``pages`` is either the
    page's text (rendered one line per ``\\n``) or ``None`` for a page that
    carries an embedded image and no extractable text at all."""
    font_num = 3
    page_nums = []
    content_nums = []
    next_num = 4
    for _ in pages:
        page_nums.append(next_num)
        next_num += 1
        content_nums.append(next_num)
        next_num += 1
    has_image = any(page is None for page in pages)
    image_num = next_num if has_image else None
    total_objects = next_num - 1 + (1 if has_image else 0)

    objects = [b""] * total_objects
    objects[0] = b"<< /Type /Catalog /Pages 2 0 R >>"
    kids = " ".join(f"{n} 0 R" for n in page_nums)
    objects[1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_nums)} >>".encode()
    objects[2] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    if image_num is not None:
        pixel = bytes([200, 30, 30])  # 1x1 red, DeviceRGB, uncompressed
        objects[image_num - 1] = (
            b"<< /Type /XObject /Subtype /Image /Width 1 /Height 1 "
            b"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Length 3 >>\n"
            b"stream\n" + pixel + b"\nendstream"
        )

    for i, page in enumerate(pages):
        page_num, content_num = page_nums[i], content_nums[i]
        resources = f"/Font << /F1 {font_num} 0 R >>"
        if page is None:
            resources += f" /XObject << /Im0 {image_num} 0 R >>"
        objects[page_num - 1] = (
            f"<< /Type /Page /Parent 2 0 R /Resources << {resources} >> "
            f"/MediaBox [0 0 200 200] /Contents {content_num} 0 R >>"
        ).encode()

        if page is None:
            stream = b"q 100 0 0 100 50 50 cm /Im0 Do Q"
        else:
            parts = ["BT", "/F1 12 Tf"]
            y = 150
            for line in page.split("\n"):
                escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(
                    ")", r"\)"
                )
                parts.append(f"1 0 0 1 20 {y} Tm")
                parts.append(f"({escaped}) Tj")
                y -= 14
            parts.append("ET")
            stream = "\n".join(parts).encode()
        objects[content_num - 1] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode()
            + stream
            + b"\nendstream"
        )

    return _pdf_from_objects(objects)


def test_normalize_converts_a_new_plain_text_unit(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "First paragraph.\n\nSecond paragraph.")
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.added_units == ["SRC-000001"]
    assert report.converted == ["SRC-000001"]
    assert report.skipped == []

    (record,) = read_all(repository)
    assert record.id == "SRC-000001"
    assert record.paragraphs == ["First paragraph.", "Second paragraph."]
    assert record.original_file == "raw/a.txt"
    assert record.converter == CONVERTERS[".txt"][1]()
    assert len(record.raw_sha256) == 64


def test_a_non_utf8_plain_text_unit_converts_instead_of_stopping_the_pass(tmp_path):
    """One Windows-1252 file must not stop the whole pass. Found on the
    Enron slice (#15): a .txt attachment carried 0x82, the UTF-8 decode
    raised, and every unit after it stayed unwritten."""
    evidence_root = tmp_path / "evidence"
    full = evidence_root / "raw" / "a.txt"
    full.parent.mkdir(parents=True)
    full.write_bytes("He said \u201cno\u201d.\n\nSecond.".encode("cp1252"))
    _write_raw_file(evidence_root, "b.txt", "After it.")
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001", "SRC-000002"]
    first, second = read_all(repository)
    assert first.paragraphs == ["He said \u201cno\u201d.", "Second."]
    assert second.paragraphs == ["After it."]


def test_a_unit_whose_converter_raises_is_reported_and_the_pass_goes_on(tmp_path):
    """A corrupt pdf attachment stopped a whole Enron normalize run (#106).
    The failed unit gets no record and is named in the report; the units
    after it are still converted."""
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.pdf", "not a pdf at all")
    _write_raw_file(evidence_root, "b.txt", "After it.")
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert list(report.failed) == ["SRC-000001"]
    assert "SRC-000001" not in report.converted
    assert report.converted == ["SRC-000002"]
    (record,) = read_all(repository)
    assert record.id == "SRC-000002"


def test_a_run_over_unchanged_input_produces_no_diff(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "Hello.\n\nWorld.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    before = record_path.read_bytes()

    report = normalize(repository, evidence_root)

    after = record_path.read_bytes()
    assert before == after
    assert report.converted == []
    assert report.skipped == ["SRC-000001"]


def test_a_new_unit_that_sorts_before_existing_ones_renumbers_nothing(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "m.txt", "Middle file.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    _write_raw_file(evidence_root, "a.txt", "Alphabetically first.")
    report = normalize(repository, evidence_root)

    records = {r.original_file: r.id for r in read_all(repository)}
    assert records["raw/m.txt"] == "SRC-000001"
    assert records["raw/a.txt"] == "SRC-000002"
    assert report.added_units == ["SRC-000002"]
    assert report.converted == ["SRC-000002"]


def test_deleting_a_raw_unit_leaves_its_number_reserved_and_skips_it(tmp_path):
    evidence_root = tmp_path / "evidence"
    path_a = _write_raw_file(evidence_root, "a.txt", "Gone soon.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    path_a.unlink()
    _write_raw_file(evidence_root, "b.txt", "New arrival.")
    report = normalize(repository, evidence_root)

    records = {r.original_file: r.id for r in read_all(repository)}
    assert records["raw/b.txt"] == "SRC-000002"
    # The deleted unit's record is untouched; its ID was never reassigned.
    assert "SRC-000001.md" in [
        p.name for p in (repository.root / NORMALIZED_RELATIVE_PATH).glob("*.md")
    ]
    assert report.added_units == ["SRC-000002"]


def test_a_changed_raw_hash_reconverts_only_that_unit(tmp_path):
    evidence_root = tmp_path / "evidence"
    path_a = _write_raw_file(evidence_root, "a.txt", "Original.")
    _write_raw_file(evidence_root, "b.txt", "Unchanged.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    path_a.write_text("Changed content.")
    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001"]
    assert report.skipped == ["SRC-000002"]
    records = {r.id: r for r in read_all(repository)}
    assert records["SRC-000001"].paragraphs == ["Changed content."]


def test_a_changed_converter_version_reconverts_the_unit(tmp_path, monkeypatch):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "Content.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    from memoria import normalize as normalize_module

    monkeypatch.setitem(
        normalize_module.CONVERTERS,
        ".txt",
        (normalize_module.convert_plain_text, lambda: "plain-text 2"),
    )
    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001"]
    (record,) = read_all(repository)
    assert record.converter == "plain-text 2"


def test_all_forces_every_unit_to_reconvert(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "One.")
    _write_raw_file(evidence_root, "b.txt", "Two.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    report = normalize(repository, evidence_root, force_all=True)

    assert sorted(report.converted) == ["SRC-000001", "SRC-000002"]
    assert report.skipped == []


def test_reconverting_with_force_all_is_still_byte_identical(tmp_path):
    """The real idempotence check: forcing a reconversion of unchanged input
    must reproduce exactly the same bytes, not merely be skipped. A
    skip-path-only byte-identity test would pass even if the converter or
    the serializer were nondeterministic, since it would never re-run
    either of them on unchanged input."""
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "One.\n\nTwo.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    before = record_path.read_bytes()

    report = normalize(repository, evidence_root, force_all=True)

    assert report.converted == ["SRC-000001"]
    assert record_path.read_bytes() == before


# --- converter pinning and the paragraph-hash drift report (#79) ----------


def test_all_with_unchanged_pins_reports_zero_changed_paragraph_hashes(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "One.\n\nTwo.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    before = record_path.read_bytes()

    report = normalize(repository, evidence_root, force_all=True)

    assert report.paragraph_drift == {}
    assert record_path.read_bytes() == before


def test_a_converter_change_reports_the_exact_count_of_changed_paragraph_hashes(
    tmp_path, monkeypatch
):
    """A fixture converter whose output differs by whitespace: two of three
    paragraphs shift, one does not, so the drift report must name exactly
    two - not three (paragraph count is unchanged) and not zero (the
    unchanged paragraph's hash really is unchanged)."""
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "First.\n\nSecond.\n\nThird.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    from memoria import normalize as normalize_module

    def _shifted(raw_bytes: bytes) -> normalize_module.ConversionDraft:
        draft = normalize_module.convert_plain_text(raw_bytes)
        paragraphs = list(draft.paragraphs)
        paragraphs[0] = paragraphs[0] + " "  # whitespace-only drift
        paragraphs[1] = paragraphs[1] + " "  # whitespace-only drift
        return replace(draft, paragraphs=paragraphs)

    monkeypatch.setitem(
        normalize_module.CONVERTERS, ".txt", (_shifted, lambda: "plain-text 2")
    )
    report = normalize(repository, evidence_root, force_all=True)

    assert report.converted == ["SRC-000001"]
    assert report.paragraph_drift == {"SRC-000001": 2}


def test_a_brand_new_unit_reports_no_paragraph_drift(tmp_path):
    """A unit with no prior record has nothing to drift against - its
    paragraphs are new content, not a converter's reconversion of old
    content."""
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "First.\n\nSecond.")
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001"]
    assert report.paragraph_drift == {}


def test_normalize_records_the_pinned_converter_version_in_the_manifest(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.txt", "Content.")
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    converters = load_converter_pins(evidence_root / DEFAULT_MANIFEST_RELATIVE_PATH)
    assert converters == {".txt": CONVERTERS[".txt"][1]()}


def test_a_converter_pin_is_preserved_when_its_suffix_disappears(tmp_path):
    """A suffix with no unit left in the corpus (the unit was deleted) keeps
    its last-recorded pin rather than losing it - #79's manifest is a
    record of what has been pinned, not only of what is present today."""
    evidence_root = tmp_path / "evidence"
    path_a = _write_raw_file(evidence_root, "a.txt", "Content.")
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)

    path_a.unlink()
    normalize(repository, evidence_root)

    converters = load_converter_pins(evidence_root / DEFAULT_MANIFEST_RELATIVE_PATH)
    assert converters == {".txt": CONVERTERS[".txt"][1]()}


def test_a_unit_with_no_registered_converter_is_reported_but_not_converted(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_file(evidence_root, "a.xlsx", "not a format memoria converts")
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.added_units == ["SRC-000001"]
    assert report.unconvertible == ["SRC-000001"]
    assert report.converted == []
    assert read_all(repository) == []


# --- email: a message is a raw unit finer than the file (#78) -------------


def test_mbox_thread_yields_three_records_with_a_shared_thread_and_a_reply_chain(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="Kickoff message.",
            ),
            _email_message(
                from_="bob@example.com", to="alice@example.com",
                date="Mon, 17 Oct 2011 10:00:00 -0500", message_id="<m2@x>",
                in_reply_to="<m1@x>", body="First reply.",
            ),
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 11:00:00 -0500", message_id="<m3@x>",
                in_reply_to="<m2@x>", body="Second reply.",
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    records = read_all(repository)
    assert len(records) == 3
    assert len({r.id for r in records}) == 3
    assert len({r.thread_id for r in records}) == 1
    assert all(r.source_type == "email" for r in records)
    by_body = {r.paragraphs[0]: r for r in records}
    assert by_body["Kickoff message."].in_reply_to == ""
    assert by_body["First reply."].in_reply_to == by_body["Kickoff message."].id
    assert by_body["Second reply."].in_reply_to == by_body["First reply."].id


def test_a_reply_quoting_its_parent_drops_the_quote_and_search_hits_the_parent_only(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="The quokka statistic belongs only to this message.",
            ),
            _email_message(
                from_="bob@example.com", to="alice@example.com",
                date="Mon, 17 Oct 2011 10:00:00 -0500", message_id="<m2@x>",
                in_reply_to="<m1@x>",
                body=(
                    "Sounds good.\n\n"
                    "On Mon, Oct 17, 2011, alice@example.com wrote:\n"
                    "> The quokka statistic belongs only to this message."
                ),
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    records = read_all(repository)
    build_index(repository, records)

    reply = next(r for r in records if r.paragraphs == ["Sounds good."])
    assert reply.quoted_excised is True

    results = search(repository, "quokka")
    assert len(results) == 1
    assert results[0].src_id != reply.id


def test_an_interleaved_reply_keeps_sender_lines_in_order_and_drops_quote_markers(tmp_path):
    evidence_root = tmp_path / "evidence"
    interleaved_body = (
        "> Original point one.\n"
        "Agreed on point one.\n"
        "> Original point two.\n"
        "Disagree on point two."
    )
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="Original point one.\n\nOriginal point two.",
            ),
            _email_message(
                from_="bob@example.com", to="alice@example.com",
                date="Mon, 17 Oct 2011 10:00:00 -0500", message_id="<m2@x>",
                in_reply_to="<m1@x>", body=interleaved_body,
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    reply = next(r for r in read_all(repository) if r.in_reply_to)
    assert reply.quoted_excised is True
    assert reply.paragraphs == ["Agreed on point one.\nDisagree on point two."]


def test_a_spreadsheet_attachment_is_listed_stored_and_gets_no_record(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="See the attached budget.",
                attachment=("budget.xlsx", b"pretend-spreadsheet-bytes"),
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    (record,) = read_all(repository)
    assert record.attachments == [{"filename": "budget.xlsx", "type": "xlsx"}]

    attachment_path = evidence_root / "raw" / "thread.mbox.attachments" / "0000-00-budget.xlsx"
    assert attachment_path.is_file()
    assert attachment_path.read_bytes() == b"pretend-spreadsheet-bytes"
    # The container (thread.mbox) and the attachment file both get reserved
    # IDs but no record of their own.
    assert len(report.unconvertible) == 2


def test_email_headers_are_in_frontmatter_and_never_in_the_body(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com", cc="carol@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="Body text only.",
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    (record,) = read_all(repository)
    assert record.email_from == "alice@example.com"
    assert record.email_to == "bob@example.com"
    assert record.email_cc == "carol@example.com"
    assert record.paragraphs == ["Body text only."]

    markdown = record_to_markdown(record)
    assert "from: alice@example.com" in markdown
    for header_value in ("alice@example.com", "bob@example.com", "carol@example.com"):
        assert header_value not in "\n".join(record.paragraphs)


def test_converting_the_same_mbox_fixture_twice_is_byte_identical(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="One.\n\nTwo.",
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    (record_path,) = (repository.root / NORMALIZED_RELATIVE_PATH).glob("SRC-*.md")
    before = record_path.read_bytes()

    report = normalize(repository, evidence_root)

    assert record_path.read_bytes() == before
    assert report.converted == []
    (skipped_record,) = read_all(repository)
    assert report.skipped == [skipped_record.id]
    assert skipped_record.converter == EMAIL_CONVERTER_VERSION


def test_normalized_email_corpus_passes_validate(tmp_path):
    """Regression: a message entry shares its `path` with the export file
    (`_process_email_containers`'s docstring), and `validate`'s per-entry
    hash check used to hash that shared path against the message's own
    (different) `sha256`, failing every email record it validated."""
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="Kickoff message.",
                attachment=("budget.xlsx", b"pretend-spreadsheet-bytes"),
            ),
            _email_message(
                from_="bob@example.com", to="alice@example.com",
                date="Mon, 17 Oct 2011 10:00:00 -0500", message_id="<m2@x>",
                in_reply_to="<m1@x>", body="First reply.",
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    assert validate(evidence_root, repo_root=repository.root) == []


def test_a_standalone_eml_file_converts_to_one_record(tmp_path):
    """A `.eml` file is an export holding exactly one message (part 05
    §5.4's "the standard library for mbox and .eml"): it goes through the
    same message-is-a-raw-unit handling as an `.mbox`, so it gets its own
    email record (and the file itself a reserved but record-less ID), with
    no parent to resolve `in_reply_to`/`thread_id` against but itself."""
    evidence_root = tmp_path / "evidence"
    eml_path = evidence_root / "raw" / "standalone.eml"
    eml_path.parent.mkdir(parents=True, exist_ok=True)
    message = _email_message(
        from_="alice@example.com", to="bob@example.com",
        date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<standalone@x>",
        body="Standalone message body.",
    )
    eml_path.write_bytes(bytes(message))
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert len(report.unconvertible) == 1  # the .eml file itself
    (record,) = read_all(repository)
    assert record.source_type == "email"
    assert record.paragraphs == ["Standalone message body."]
    assert record.in_reply_to == ""
    assert record.thread_id == "standalone@x"


def test_deleting_an_mbox_marks_its_message_ids_deleted_instead_of_dropping_them(tmp_path):
    """Regression: `_process_email_containers` skipped a deleted container
    outright, so its message sub-entries (told apart only by
    `email_message_index`, sharing the container's own `path`) fell out of
    `combined` entirely instead of being carried forward as `deleted` - the
    rule ADR-0006 and manifest.py's docstring require ("Nothing is reused").
    Reissuing their numbers would silently repoint any existing citation
    into that ID range at different evidence."""
    evidence_root = tmp_path / "evidence"
    _write_mbox(
        evidence_root,
        "thread.mbox",
        [
            _email_message(
                from_="alice@example.com", to="bob@example.com",
                date="Mon, 17 Oct 2011 09:00:00 -0500", message_id="<m1@x>",
                body="Kickoff message.",
            ),
            _email_message(
                from_="bob@example.com", to="alice@example.com",
                date="Mon, 17 Oct 2011 10:00:00 -0500", message_id="<m2@x>",
                in_reply_to="<m1@x>", body="First reply.",
            ),
        ],
    )
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    manifest_path = evidence_root / DEFAULT_MANIFEST_RELATIVE_PATH
    message_ids = sorted(
        e.id for e in load_manifest(manifest_path) if "email_message_index" in e.extra
    )
    assert len(message_ids) == 2

    (evidence_root / "raw" / "thread.mbox").unlink()
    normalize(repository, evidence_root)

    entries_by_id = {e.id: e for e in load_manifest(manifest_path)}
    for message_id in message_ids:
        assert entries_by_id[message_id].deleted is True

    _write_raw_file(evidence_root, "new.txt", "A brand new unit.")
    report = normalize(repository, evidence_root)

    assert report.added_units == ["SRC-000004"]  # not a reissued SRC-000003

# --- docx (#77) --------------------------------------------------------


def test_normalize_converts_a_docx_with_full_structure(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(evidence_root, "a.docx", _make_docx())
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001"]
    (record,) = read_all(repository)
    body = "\n".join(record.paragraphs)
    assert "# A Sample Heading" in body
    assert "**bold**" in body and "*italic*" in body
    assert "* First bullet" in body and "* Second bullet" in body
    assert "| Header A | Header B |" in body
    assert record.images == ["image1.png"]
    assert record.converter == CONVERTERS[".docx"][1]()
    # Not embedded: no image data anywhere in the body (docs/
    # normalized-record-schema.md, "images").
    assert "data:image" not in body


def test_normalize_a_docx_with_no_images_reports_an_empty_list(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(evidence_root, "a.docx", _make_docx(with_image=False))
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    (record,) = read_all(repository)
    assert record.images == []


def test_docx_table_structure_is_kept_verbatim(tmp_path):
    """Whitespace policy: a paragraph's own line structure is never
    reflowed - the table's four rows stay four lines, not run together."""
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(evidence_root, "a.docx", _make_docx())
    repository = Repository(root=tmp_path / "repo")

    normalize(repository, evidence_root)

    (record,) = read_all(repository)
    table = next(p for p in record.paragraphs if p.startswith("|"))
    assert table.count("\n") == 3


def test_reconverting_a_docx_with_force_all_is_byte_identical(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(evidence_root, "a.docx", _make_docx())
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    before = record_path.read_bytes()

    report = normalize(repository, evidence_root, force_all=True)

    assert report.converted == ["SRC-000001"]
    assert record_path.read_bytes() == before


# --- pdf (#77) -----------------------------------------------------------


def test_normalize_converts_a_two_page_pdf_with_a_page_marker(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(
        evidence_root, "a.pdf", _make_pdf(["Page one text.", "Page two text."])
    )
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001"]
    (record,) = read_all(repository)
    assert record.paragraphs == [
        "Page one text.",
        "<!-- page 2 -->",
        "Page two text.",
    ]
    assert real_paragraphs(record) == ["Page one text.", "Page two text."]

    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    markdown = record_path.read_text(encoding="utf-8")
    assert '<a id="src-000001-p1"></a>' in markdown
    assert '<a id="src-000001-p2"></a>' in markdown
    # The marker itself never earns an anchor of its own.
    assert '<a id="src-000001-p3"></a>' not in markdown
    assert record.converter == CONVERTERS[".pdf"][1]()


def test_pdf_page_marker_is_absent_from_the_fts5_index(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(
        evidence_root, "a.pdf", _make_pdf(["Page one text.", "Page two text."])
    )
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    records = read_all(repository)

    build_index(repository, records)

    con = sqlite3.connect(repository.root / INDEX_RELATIVE_PATH)
    rows = con.execute("SELECT anchor, text FROM records ORDER BY anchor").fetchall()
    con.close()
    assert rows == [
        ("src-000001-p1", "Page one text."),
        ("src-000001-p2", "Page two text."),
    ]


def test_normalize_an_image_only_pdf_produces_a_stub_record(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(evidence_root, "a.pdf", _make_pdf([None]))
    repository = Repository(root=tmp_path / "repo")

    report = normalize(repository, evidence_root)

    assert report.converted == ["SRC-000001"]
    (record,) = read_all(repository)
    assert record.paragraphs == []

    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    markdown = record_path.read_text(encoding="utf-8")
    assert "<a id=" not in markdown


def test_reconverting_a_pdf_with_force_all_is_byte_identical(tmp_path):
    evidence_root = tmp_path / "evidence"
    _write_raw_binary_file(
        evidence_root, "a.pdf", _make_pdf(["Page one text.", "Page two text."])
    )
    repository = Repository(root=tmp_path / "repo")
    normalize(repository, evidence_root)
    record_path = repository.root / NORMALIZED_RELATIVE_PATH / "SRC-000001.md"
    before = record_path.read_bytes()

    report = normalize(repository, evidence_root, force_all=True)

    assert report.converted == ["SRC-000001"]
    assert record_path.read_bytes() == before
